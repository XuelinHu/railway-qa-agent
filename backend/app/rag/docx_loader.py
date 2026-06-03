from dataclasses import dataclass
from pathlib import Path

from docx import Document


@dataclass(frozen=True)
class DocumentBlock:
    text: str
    block_type: str
    heading: str | None
    index: int


@dataclass(frozen=True)
class DocumentChunk:
    text: str
    source_file: str
    heading: str | None
    chunk_index: int


def load_docx_blocks(path: Path) -> list[DocumentBlock]:
    document = Document(path)
    blocks: list[DocumentBlock] = []
    current_heading: str | None = None

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue

        style_name = (
            paragraph.style.name.lower()
            if paragraph.style and paragraph.style.name
            else ""
        )
        looks_like_heading = "heading" in style_name or text.startswith(("第", "Chapter"))
        block_type = "heading" if looks_like_heading else "paragraph"
        if block_type == "heading" and len(text) <= 120:
            current_heading = text

        blocks.append(
            DocumentBlock(
                text=text,
                block_type=block_type,
                heading=current_heading,
                index=len(blocks),
            )
        )

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if not cells:
                continue
            blocks.append(
                DocumentBlock(
                    text=" | ".join(cells),
                    block_type="table_row",
                    heading=current_heading,
                    index=len(blocks),
                )
            )

    return blocks


def chunk_blocks(
    blocks: list[DocumentBlock],
    source_file: str,
    max_chars: int = 1200,
    overlap_chars: int = 160,
) -> list[DocumentChunk]:
    chunks: list[DocumentChunk] = []
    buffer: list[str] = []
    current_heading: str | None = None
    current_len = 0

    for block in blocks:
        next_len = len(block.text) + 1
        if buffer and current_len + next_len > max_chars:
            chunks.append(
                DocumentChunk(
                    text="\n".join(buffer),
                    source_file=source_file,
                    heading=current_heading,
                    chunk_index=len(chunks),
                )
            )
            overlap = "\n".join(buffer)[-overlap_chars:]
            buffer = [overlap] if overlap else []
            current_len = len(overlap)

        current_heading = block.heading or current_heading
        buffer.append(block.text)
        current_len += next_len

    if buffer:
        chunks.append(
            DocumentChunk(
                text="\n".join(buffer),
                source_file=source_file,
                heading=current_heading,
                chunk_index=len(chunks),
            )
        )

    return chunks


def normalize_text(text: str) -> str:
    return " ".join(text.replace("\u3000", " ").split())
