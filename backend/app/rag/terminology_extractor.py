import re
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from app.rag.docx_loader import normalize_text

CJK_PATTERN = re.compile(r"[\u4e00-\u9fff]")
PAIR_PATTERN = re.compile(
    r"(?P<zh>[\u4e00-\u9fff][\u4e00-\u9fffA-Za-z0-9（）()、/·．.\-\s]{0,40}?)"
    r"\s+"
    r"(?P<en>[^\u4e00-\u9fff]+?)"
    r"(?=\s+[\u4e00-\u9fff]|$)"
)


@dataclass(frozen=True)
class ExtractedTerm:
    source_term: str
    target_term: str
    source_language: str
    target_language: str
    category: str | None
    source_file: str


def extract_terms_from_docx(path: Path, limit: int | None = None) -> list[ExtractedTerm]:
    document = Document(path)
    category: str | None = None
    terms: list[ExtractedTerm] = []
    seen: set[tuple[str, str]] = set()

    def add_pairs(line: str) -> None:
        nonlocal terms
        for source_term, target_term in extract_pairs_from_line(line):
            key = (source_term.casefold(), target_term.casefold())
            if key in seen:
                continue
            seen.add(key)
            terms.append(
                ExtractedTerm(
                    source_term=source_term,
                    target_term=target_term,
                    source_language="zh",
                    target_language="en",
                    category=category,
                    source_file=path.name,
                )
            )
            if limit and len(terms) >= limit:
                return

    for paragraph in document.paragraphs:
        line = normalize_text(paragraph.text)
        if not line:
            continue
        if is_category_line(line):
            category = line[:120]
            continue
        add_pairs(line)
        if limit and len(terms) >= limit:
            return terms

    for table in document.tables:
        for row in table.rows:
            raw_cells = [cell.text for cell in row.cells]
            add_pairs_from_cells(raw_cells, add_pairs)
            if limit and len(terms) >= limit:
                return terms

    return terms


def add_pairs_from_cells(raw_cells: list[str], add_pairs: Callable[[str], None]) -> None:
    cells = [cell for cell in raw_cells if normalize_text(cell)]
    if len(cells) >= 2:
        zh_lines = split_cell_lines(cells[0])
        en_lines = split_cell_lines(cells[1])
        if len(zh_lines) == len(en_lines) and len(zh_lines) > 1:
            for source_term, target_term in zip(zh_lines, en_lines, strict=True):
                if CJK_PATTERN.search(source_term) and not CJK_PATTERN.search(target_term):
                    add_pairs(f"{source_term} {target_term}")
            return

    add_pairs(normalize_text(" ".join(cells)))


def split_cell_lines(text: str) -> list[str]:
    return [normalize_text(line) for line in text.splitlines() if normalize_text(line)]


def extract_pairs_from_line(line: str) -> list[tuple[str, str]]:
    if should_skip_line(line):
        return []

    pairs: list[tuple[str, str]] = []
    for match in PAIR_PATTERN.finditer(line):
        source_term = normalize_source_term(match.group("zh"))
        target_term = normalize_target_term(match.group("en"))
        if is_valid_pair(source_term, target_term):
            pairs.append((source_term, target_term))
    return pairs


def normalize_source_term(term: str) -> str:
    term = normalize_text(term)
    term = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", term)
    return term.strip(" ,;:，；：")


def normalize_target_term(term: str) -> str:
    term = normalize_text(term)
    return term.strip(" ,;:，；：")


def is_category_line(line: str) -> bool:
    return (
        len(line) <= 80
        and CJK_PATTERN.search(line) is not None
        and any(token in line.lower() for token in ("vocabulary", "词汇"))
    )


def should_skip_line(line: str) -> bool:
    lowered = line.lower()
    skip_tokens = (
        "dictionary",
        "publishing",
        "出版社",
        "主编",
        "目录",
        "copyright",
    )
    return any(token in lowered for token in skip_tokens)


def is_valid_pair(source_term: str, target_term: str) -> bool:
    if not source_term or not target_term:
        return False
    if len(source_term) > 50 or len(target_term) > 120:
        return False
    if not CJK_PATTERN.search(source_term):
        return False
    if CJK_PATTERN.search(target_term):
        return False
    if not re.search(r"[A-Za-z]", target_term):
        return False
    return True
